import mysql.connector as mysql
from docx.api import Document
import pprint

table_names = ["Item_Table", "Customer_Table", "Salesman_Table", "Sales_Order_Table", "Customer_Table1", "Loan_Table", "Borrower_Table", "Account_Table", "Depositor_Table"]
document = Document("/home/avi/Downloads/file.docx")

# CREATING A DATABSE NAMED STORE
database_name = "store1"
db = mysql.connect(
	host = "localhost", 
	user = "root", 
	passwd = "",
	)
cursor = db.cursor()
try:
	cursor.execute("CREATE DATABASE " + database_name)
	cursor.execute("USE " + database_name)
except:
	pass
db.commit()
###############################################################

decimal = ["item_rate", "Sell_prize", "Bal_due"]
integer = ["Pin", "Bal_due", "Salary", "Tot_quiantity_order", "Amount", "Balance"]

sqlFile = open("sqlFile.txt",'w+')
# USING THIS LITTLE CODE TO DELETE ALL THE TABLES, IN CASE WE WANT TO RE-RUN THE CODE ON THE SAME DATABASE WITH THE SAME TABLES
for i in range(len(document.tables)):
	db = mysql.connect(
		host = "localhost", 
		user = "root", 
		passwd = "",
		database = database_name
		)
	cursor = db.cursor()
	try:
		cursor.execute("DROP TABLE " + table_names[i]);
	except:
		pass
	db.commit();
###############################################################

for i in range(len(document.tables)):
	# MAKING A LIST OF ALL THE DATA IN A TABLE
	table = document.tables[i]
	data = []
	keys = None
	for j, row in enumerate(table.rows):
		text = (cell.text for cell in row.cells)
		if(j==0):
			keys = tuple(text)
			continue
		row_data = tuple(text)
		data.append(row_data)

	# pprint.pprint(data)
	###############################################################

	# CREATING AND INSERTING DATA INTO A TABLE
	db = mysql.connect(
		host = "localhost", 
		user = "root", 
		passwd = "Maninhell1#",
		database = database_name
		)

	cursor = db.cursor()
	create_query = "CREATE TABLE " + table_names[i] + " ("
	sqlFile.write(str(i+1)+". " +table_names[i].upper()+'\n')
	for j in range(len(keys)):
		item = str(keys[j]).strip().replace(" ","").replace("\n","")
		if item in decimal:
			item += " FLOAT(2) NOT NULL, "
		elif item in integer:
			item += " INT NOT NULL, "
		else:
			item +=" VARCHAR(30) NOT NULL, "
		create_query += item
		if(j == len(keys)-1):
			create_query = create_query[:len(create_query)-2] + ");"

	# print(create_query)
	sqlFile.write(create_query + "\n")
	cursor.execute(create_query)

	insert_query = "INSERT INTO " + table_names[i] + " VALUES ("
	for j in range(len(data)):
		row_data = data[j]
		item = ""
		for k in row_data:
			item += "'" + str(k).strip().replace(" ","").replace("\n","") + "', "
		item = item[:-2] + ');'
		# print(insert_query + item)
		sqlFile.write(insert_query + item + "\n")
		cursor.execute(insert_query + item)

	db.commit()
	###############################################################
	sqlFile.write("Table " + str(i+1) + " complete.\n\n")
	print("Table " + str(i+1) + " complete.")

sqlFile.close()
